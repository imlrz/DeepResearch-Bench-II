# -*- coding: utf-8 -*-

import os
import re
import json
import argparse
import builtins
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Tuple, Optional

from gemini_client import GeminiClient, GeminiInput, GeminiOutput, get_config

# =========================
# Global variables
# =========================
# Client instance (initialized in main)
client: Optional[GeminiClient] = None


def get_default_config():
    """Get default configuration values from the .env file (with defaults)."""
    return {
        'pdf_dir': get_config('PDF_DIR', 'report'),
        'out_jsonl': get_config('OUT_JSONL', 'result.jsonl'),
        'tasks_jsonl': get_config('TASKS_JSONL', 'tasks_and_rubrics.jsonl'),
        'chunk_size': int(get_config('CHUNK_SIZE', '50')),
        'max_workers': int(get_config('MAX_WORKERS', '10')),
        'max_retries': int(get_config('MAX_RETRIES', '10')),
        'max_paper_chars': int(get_config('MAX_PAPER_CHARS', '150000')),
        'log_file': get_config('LOG_FILE', 'run_evaluation.log'),
    }


def setup_print_logger(log_file: str):
    """
    Simple logging helper: mirror all print output into a log file.

    - Does not change existing usage of print;
    - Console output still works as usual;
    - Each printed line will be appended to log_file.
    """
    if not log_file:
        return

    log_dir = os.path.dirname(log_file)
    if log_dir:
        os.makedirs(log_dir, exist_ok=True)

    # Avoid wrapping print multiple times
    if getattr(builtins, "_orig_print", None) is None:
        builtins._orig_print = builtins.print

    def logged_print(*args, **kwargs):
        # First print to console
        builtins._orig_print(*args, **kwargs)
        try:
            text = " ".join(str(a) for a in args)
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(text + "\n")
        except Exception:
            # Logging failures must not break main flow
            pass

    builtins.print = logged_print

# =========================
# Prompt template (three-way classification)
# =========================
PROMPT_TEMPLATE = """
You will receive an article, a task, and a list of grading rubric items. Your job is to assess whether the article satisfies each rubric item, and provide a THREE-WAY score for EACH rubric item.

Scoring rule per rubric item (strict):
- Score = 1: The article clearly satisfies the rubric item AND the specific supporting sentence(s) do NOT cite any reference listed in "blocked" (match by title/urls). For numerical data, exact values must be explicitly listed and match the rubric.
- Score = 0: The article does NOT mention this rubric item at all.
- Score = -1: The article mentions this rubric item, BUT the supporting sentence(s) cite a blocked reference.

For EACH rubric item, you MUST provide:
1. "score": 1, 0, or -1
2. "reason": A brief explanation
3. "evidence": The specific supporting sentence(s) from the article (empty string if score is 0)

The input format is:
<input_format>
{{
    "task": "...",
    "rubric_items": ["rubric item 1", "rubric item 2", ...],
    "blocked": {{
        "title": "...",
        "authors": ["...", "..."],
        "urls": ["...", "..."]
    }}
}}
</input_format>

Your output MUST strictly follow this JSON format (no extra keys, and the rubric item text MUST match the input EXACTLY):
<output_format>
{{
    "results": [
        {{
            "rubric_item": "rubric item 1",
            "score": 1 or 0 or -1,
            "reason": "brief explanation",
            "evidence": "supporting sentence(s) from the article"
        }},
        {{
            "rubric_item": "rubric item 2",
            "score": 1 or 0 or -1,
            "reason": "brief explanation",
            "evidence": "supporting sentence(s) from the article"
        }},
        ...
    ]
}}
</output_format>

CRITICAL: You MUST return results for ALL rubric items in the input, and the "rubric_item" text MUST match the input text EXACTLY (character-level match).

<passage>
{paper}
</passage>
<task_and_rubric>
{rubric}
</task_and_rubric>
Now, please begin your generation
"""

PAPER_PLACEHOLDER = "(The PDF is attached to this message; please treat its full content as the <passage> text when searching and editing.)"

# =========================
# Load tasks.jsonl index
# =========================
def load_tasks_data(path: str):
    tasks_data = {}
    original_data = {}
    if not os.path.exists(path):
        print(f"[warn] tasks file not found: {path}, skip all items by default")
        return tasks_data, original_data
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            if not line.strip():
                continue
            try:
                obj = json.loads(line)
                idx_val = obj.get("idx", None)
                content = obj.get("content", "")
                if idx_val is not None:
                    try:
                        key = int(idx_val)
                    except Exception:
                        continue
                    tasks_data[key] = content
                    original_data[key] = obj
            except Exception:
                continue
    return tasks_data, original_data

# Lazy load (populated in main)
TASKS_DATA = {}
ORIGINAL_DATA = {}

# =========================
# Document content extraction
# =========================
def _extract_docx_content(path: str) -> Tuple[str, List[Tuple[str, bytes]]]:
    """
    Extract content from a .docx file.

    Returns:
        (text_content, [(image_mime, image_bytes)]).
    """
    try:
        from docx import Document  # python-docx
    except Exception as e:
        print(f"[warn] python-docx is not installed, cannot parse .docx: {e}")
        return "", []
    try:
        doc = Document(path)
    except Exception as e:
        print(f"[warn] failed to read .docx file: {e}")
        return "", []
    
    # Extract all contents (paragraphs + tables, in document order)
    all_content = []
    images = []
    
    # Extract all images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_blob = rel.target_part.blob
                # Get content_type from the relationship
                content_type = rel.target_part.content_type
                images.append((content_type, image_blob))
            except Exception as e:
                print(f"[warn] failed to extract image: {e}")
    
    # Iterate over all elements in the document body (paragraphs and tables)
    for element in doc.element.body:
        # Paragraphs
        if element.tag.endswith('p'):
            # Find the corresponding Paragraph object
            for p in doc.paragraphs:
                if p._element == element:
                    text = (p.text or "").strip()
                    if text:
                        all_content.append(text)
                    break
        # Tables
        elif element.tag.endswith('tbl'):
            # Find the corresponding Table object
            for table in doc.tables:
                if table._element == element:
                    # Extract table content and convert to Markdown format
                    table_text = _table_to_markdown(table)
                    if table_text:
                        all_content.append(table_text)
                    break
    
    content = "\n\n".join(all_content)
    
    return content, images

def _table_to_markdown(table) -> str:
    """
    Convert a docx table into Markdown formatted text.
    """
    if not table.rows:
        return ""
    
    lines = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        # Use | to separate cells
        lines.append("| " + " | ".join(cells) + " |")
        # Add separator line after the first row
        if i == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    
    return "\n".join(lines)



# =========================
# JSON fenced block extraction / parsing
# =========================
FENCED_JSON_PATTERN = r'```json\s*(.*)```'

def _try_clean_and_load(s: str):
    json_clean = re.sub(
        r'"(?P<k>.*?)"(?=\s*:)',
        lambda m: '"' + re.sub(r'(?<!\\)"', r'\"', m.group('k')) + '"',
        s
    )
    return json.loads(json_clean.strip())

def parse_model_text(text: str):
    matches = re.findall(FENCED_JSON_PATTERN, text, re.DOTALL)
    if matches:
        try:
            return _try_clean_and_load(matches[0]), True
        except json.JSONDecodeError as e:
            print(f"[warn] failed to parse fenced JSON: {e}; trying full text...")
    try:
        return _try_clean_and_load(text), True
    except json.JSONDecodeError as e:
        print(f"[warn] failed to parse JSON from full text: {e}")
        return None, False

# =========================
# Batched evaluation and validation
# =========================
def validate_batch_result(rubric_items: List[str], parsed_result: Dict) -> bool:
    """
    Validate that the model output contains all rubric_items with exact text match.
    """
    if not isinstance(parsed_result, dict):
        return False
    results = parsed_result.get("results", [])
    if not isinstance(results, list):
        return False
    if len(results) != len(rubric_items):
        return False
    
    # 检查每个 rubric_item 是否严格匹配
    returned_items = [r.get("rubric_item", "") for r in results]
    for expected in rubric_items:
        # Ensure every rubric_item from the input is present in the results
        if expected not in returned_items:
            return False
    
    return True

def query_rubric_batch(rubric_items: List[str], task: str, blocked: Dict, 
                       paper_content: str, pdf_path: str = None, 
                       extra_images: List[Tuple[str, bytes]] = None,
                       max_retries: int = 5) -> Tuple[Optional[List[Dict]], Dict]:
    """
    Query a batch of rubric items.

    Returns:
        (results_list, usage_metadata).
    Retries up to max_retries times.
    """
    global client
    if client is None:
        raise RuntimeError("GeminiClient is not initialized")
    
    rubric_input = {
        "task": task,
        "rubric_items": rubric_items,
        "blocked": blocked
    }
    rubric_json = json.dumps(rubric_input, ensure_ascii=False, indent=2)
    
    for attempt in range(max_retries):
        try:
            # Build input
            if pdf_path:
                # With file attachment
                prompt = PROMPT_TEMPLATE.format(paper=PAPER_PLACEHOLDER, rubric=rubric_json)
                input_data = GeminiInput(
                    text=prompt,
                    file_path=pdf_path,
                    extra_images=extra_images
                )
            else:
                # Text-only
                prompt = PROMPT_TEMPLATE.format(paper=paper_content, rubric=rubric_json)
                input_data = GeminiInput(
                    text=prompt,
                    extra_images=extra_images
                )
            
            # Call the client
            output = client.query(input_data)
            
            if not output.text:
                print(f"[warn] batch attempt {attempt+1}/{max_retries} returned empty text")
                continue
            
            # Parse JSON
            parsed, ok = parse_model_text(output.text)
            if not ok:
                print(f"[warn] batch attempt {attempt+1}/{max_retries} JSON parse failed")
                continue
            
            # Validate result
            if not validate_batch_result(rubric_items, parsed):
                print(f"[warn] batch attempt {attempt+1}/{max_retries} validation failed: rubric_item mismatch or wrong count")
                continue
            
            # Success
            return parsed["results"], output.usage_metadata
            
        except Exception as e:
            print(f"[warn] batch attempt {attempt+1}/{max_retries} request error: {e}")
            continue
    
    # All attempts failed
    return None, {}

# =========================
# Per-file processing (batched version)
# =========================
def process_one_with_chunking(idx: int, pdf_path: str, rubric_content: Dict, chunk_size: int = 0, max_paper_chars: int = 150000, max_retries: int = 5):
    """
    Process a single file, with batched evaluation (Gemini multimodal).

    Returns:
        (idx, result_dict, total_tokens)
    where result_dict contains scores (organized by dimension) and total_tokens.
    """
    print(f"[run] processing idx={idx}, file={os.path.basename(pdf_path)}, chunk_size={chunk_size}")
    
    # Parse rubric_content
    if not isinstance(rubric_content, dict):
        print(f"[err] idx={idx} rubric_content has invalid format")
        return idx, {"error": "invalid rubric_content"}, 0
    
    task = rubric_content.get("task", "")
    rubric = rubric_content.get("rubric", {})
    blocked = rubric_content.get("blocked", {})
    
    # Collect all rubric items (across all dimensions)
    all_items = []
    dimension_map = {}  # rubric_item -> dimension
    for dim in ["info_recall", "analysis", "presentation"]:
        items = rubric.get(dim, [])
        if isinstance(items, list):
            for item in items:
                all_items.append(item)
                dimension_map[item] = dim
    
    if not all_items:
        print(f"[warn] idx={idx} has no rubric items")
        return idx, {"error": "no rubric items"}, 0
    
    # Prepare document content (multimodal strategy)
    text_content = ""
    file_to_upload = None
    extra_images = []
    lower_name = pdf_path.lower()
    
    if lower_name.endswith('.pdf'):
        # For PDF, upload the file as an attachment by default
        file_to_upload = pdf_path
        print(f"[info] idx={idx} PDF will be uploaded as an attachment")
    elif lower_name.endswith('.docx'):
        # For DOCX, extract text, tables and images
        text_content, extra_images = _extract_docx_content(pdf_path)
        print(f"[info] idx={idx} DOCX extracted: text_length={len(text_content)}, image_count={len(extra_images)}")
        if not text_content and not extra_images:
            # If extraction fails, fall back to plain text
            print(f"[warn] idx={idx} DOCX extraction failed, trying as plain text")
    elif lower_name.endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp', '.tiff')):
        # Image file
        file_to_upload = pdf_path
        print(f"[info] idx={idx} image file will be uploaded as an attachment")
    elif lower_name.endswith(('.txt', '.md', '.html')):
        # Plain text file
        try:
            with open(pdf_path, 'r', encoding='utf-8', errors='ignore') as rf:
                text_content = rf.read()
            print(f"[info] idx={idx} text file read: length={len(text_content)}")
        except Exception as e:
            print(f"[warn] idx={idx} failed to read text file: {e}")
            text_content = ""
    else:
        # Unknown format, best-effort text read
        print(f"[warn] idx={idx} unknown format, trying to read as text")
        try:
            with open(pdf_path, 'r', encoding='utf-8', errors='ignore') as rf:
                text_content = rf.read()
        except Exception:
            text_content = ""
    
    # Truncate overly long text
    if text_content and len(text_content) > max_paper_chars:
        print(f"[info] idx={idx} text too long ({len(text_content)}), truncating to {max_paper_chars}")
        text_content = text_content[:max_paper_chars]
    
    # Batched processing
    if chunk_size <= 0 or chunk_size >= len(all_items):
        # No batching, process all at once
        batches = [all_items]
    else:
        # Split into batches
        batches = [all_items[i:i+chunk_size] for i in range(0, len(all_items), chunk_size)]
    
    print(f"[info] idx={idx} has {len(all_items)} rubric items, split into {len(batches)} batches")
    
    # 累积结果
    all_results = []
    total_input_tokens = 0
    total_output_tokens = 0
    total_thoughts_tokens = 0
    total_tokens_sum = 0
    all_usage_metadata = []  # save full usageMetadata for each batch
    
    for batch_idx, batch_items in enumerate(batches):
        print(f"[info] idx={idx} processing batch {batch_idx+1}/{len(batches)} ({len(batch_items)} items)")
        
        results, usage_metadata = query_rubric_batch(
            batch_items, task, blocked, text_content, 
            file_to_upload, 
            extra_images,
            max_retries
        )
        
        if results is None:
            print(f"[err] idx={idx} batch {batch_idx+1} failed")
            return idx, {"error": f"batch {batch_idx+1} failed after {max_retries} retries"}, 0
        
        all_results.extend(results)
        all_usage_metadata.append(usage_metadata)
        
        # Accumulate token counts
        total_input_tokens += usage_metadata.get("promptTokenCount", 0)
        total_output_tokens += usage_metadata.get("candidatesTokenCount", 0)
        total_thoughts_tokens += usage_metadata.get("thoughtsTokenCount", 0)
        total_tokens_sum += usage_metadata.get("totalTokenCount", 0)
    
    # Re-organize results by dimension
    scores_by_dimension = {
        "info_recall": {},
        "analysis": {},
        "presentation": {}
    }
    
    for result in all_results:
        rubric_item = result.get("rubric_item", "")
        score = result.get("score", 0)
        reason = result.get("reason", "")
        evidence = result.get("evidence", "")
        
        dim = dimension_map.get(rubric_item)
        if dim:
            scores_by_dimension[dim][rubric_item] = {
                "score": score,
                "reason": reason,
                "evidence": evidence
            }
    
    result_dict = {
        "task": task,
        "scores": scores_by_dimension,
        "usage_summary": {
            "total_tokens": total_tokens_sum,
            "input_tokens": total_input_tokens,
            "output_tokens": total_output_tokens,
            "thoughts_tokens": total_thoughts_tokens
        },
        "usage_metadata_per_batch": all_usage_metadata  # save full metadata for each batch
    }
    
    print(f"[ok] idx={idx} finished, total tokens={total_tokens_sum} (input={total_input_tokens}, output={total_output_tokens}, thoughts={total_thoughts_tokens})")
    return idx, result_dict, total_tokens_sum

# =========================
# Main entrypoint
# =========================
def main():
    global client, TASKS_DATA, ORIGINAL_DATA
    
    # Load default config
    config = get_default_config()
    
    parser = argparse.ArgumentParser(description="batched evaluation")
    parser.add_argument("--pdf_dir", default=config['pdf_dir'], help="Input directory (containing per-model subdirectories).")
    parser.add_argument("--out_jsonl", default=config['out_jsonl'], help="Output JSONL file.")
    parser.add_argument("--tasks_jsonl", default=config['tasks_jsonl'], help="Tasks and rubrics JSONL file.")
    parser.add_argument("--chunk_size", type=int, default=config['chunk_size'], help="Batch size for rubric items (0 = no batching).")
    parser.add_argument("--max_workers", type=int, default=config['max_workers'], help="Number of concurrent workers.")
    parser.add_argument("--max_retries", type=int, default=config['max_retries'], help="Maximum retry count per batch.")
    parser.add_argument("--max_paper_chars", type=int, default=config['max_paper_chars'], help="Maximum number of characters to keep from each document.")
    parser.add_argument("--log_file", default=config['log_file'], help="Log file path (all console output will also be written here).")
    parser.add_argument("--model", default=None, help="Model name (optional, overrides .env / environment).")
    parser.add_argument("--api_url", default=None, help="API URL (optional, overrides .env / environment).")
    parser.add_argument("--token", default=None, help="API token (optional, overrides .env / environment).")
    parser.add_argument("--req_id", default=None, help="Request identifier (optional, overrides .env / environment).")
    args = parser.parse_args()
    
    # Initialize print logger: mirror all prints into the log file
    setup_print_logger(args.log_file)
    print(f"[init] log file: {args.log_file}")
    
    # Load tasks data
    TASKS_DATA, ORIGINAL_DATA = load_tasks_data(args.tasks_jsonl)
    
    # Initialize Gemini client (from .env or CLI overrides)
    try:
        client = GeminiClient(
            api_url=args.api_url,
            api_token=args.token,
            model=args.model,
            request_id=args.req_id,
            verbose=True
        )
        print(f"[init] Gemini client initialized")
        print(f"  - model: {client.model}")
        print(f"  - API URL: {client.api_url}")
        print(f"  - request ID: {client.request_id}")
    except ValueError as e:
        print(f"[err] failed to initialize client: {e}")
        print(f"[hint] Please create a .env file and configure GEMINI_API_URL / GEMINI_API_TOKEN / GEMINI_MODEL, or pass them via CLI arguments.")
        return
    
    # Use local variables for convenience
    pdf_dir = args.pdf_dir
    out_jsonl = args.out_jsonl
    chunk_size = args.chunk_size
    max_retries = args.max_retries
    max_paper_chars = args.max_paper_chars
    
    if not TASKS_DATA:
        print(f"[info] no task data found in {args.tasks_jsonl}")
        return
    
    def _parse_idx_from_filename(name: str):
        m = re.match(r'^idx-(\d+)(?:\..+)?$', name, re.IGNORECASE)
        return int(m.group(1)) if m else None
    
    pairs = []  # (model, idx, fpath, content)
    if not os.path.isdir(pdf_dir):
        print(f"[info] input directory does not exist: {pdf_dir}")
        return
    
    # Traverse first-level subdirectories; subdirectory name is the model name
    for entry in os.listdir(pdf_dir):
        model_dir = os.path.join(pdf_dir, entry)
        if not os.path.isdir(model_dir):
            continue
        model_name = entry
        for fname in os.listdir(model_dir):
            idx_val = _parse_idx_from_filename(fname)
            if idx_val is None:
                continue
            fpath = os.path.join(model_dir, fname)
            content = TASKS_DATA.get(idx_val)
            if content is not None and os.path.exists(fpath):
                pairs.append((model_name, idx_val, fpath, content))
            else:
                if content is None:
                    print(f"[skip] idx={idx_val} not found in {args.tasks_jsonl} (model={model_name})")
                if not os.path.exists(fpath):
                    print(f"[skip] file does not exist: {fpath}")
    
    if not pairs:
        print("[info] no valid (model, idx, file) pairs found; nothing to do")
        return
    
    print(f"[info] found {len(pairs)} valid pairs, start parallel processing (chunk_size={chunk_size})...")
    
    # Store results
    results = {}
    
    with ThreadPoolExecutor(max_workers=min(args.max_workers, len(pairs))) as pool:
        future_to_meta = {}
        for (model, idx, pdf_path, content) in pairs:
            fut = pool.submit(process_one_with_chunking, idx, pdf_path, content, chunk_size, max_paper_chars, max_retries)
            future_to_meta[fut] = (model, idx)
        for fut in as_completed(future_to_meta.keys()):
            model, _orig_idx = future_to_meta[fut]
            ridx, result_dict, total_tokens = fut.result()
            # Bucket results by model to avoid clashes when different models share the same idx
            results.setdefault(model, {})[ridx] = result_dict
    
    # Append results to JSONL
    print(f"[info] appending results to {out_jsonl}...")
    out_dir = os.path.dirname(out_jsonl)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    with open(out_jsonl, "a", encoding="utf-8") as f:
        for model, idx_to_content in results.items():
            for idx, result_dict in idx_to_content.items():
                line_obj = {"model": model, "idx": idx, "result": result_dict}
                f.write(json.dumps(line_obj, ensure_ascii=False) + "\n")
    
    print(f"[done] processing completed, results appended to {out_jsonl}")

if __name__ == "__main__":
    main()
